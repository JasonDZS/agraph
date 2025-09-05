"""Microsoft Word document processor implementation.

This module provides functionality for extracting text and metadata from Microsoft Word
documents, supporting both modern .docx format and legacy .doc format. It handles
document content, tables, and comprehensive metadata extraction.
"""

from pathlib import Path
from typing import Any, Dict, List, Union

from .base import DocumentProcessor, ProcessingError


class WordProcessor(DocumentProcessor):
    """Document processor for Microsoft Word documents.

    This processor handles both .docx and .doc formats using different approaches:
    - .docx files: Uses python-docx library for comprehensive processing
    - .doc files: Uses docx2txt library or antiword as fallback

    Features:
    - Text extraction from paragraphs and tables
    - Metadata extraction from document properties
    - Support for both modern and legacy Word formats
    - Structured content preservation

    Dependencies:
        python-docx: Required for .docx file processing
        docx2txt: Optional, for .doc file processing
        antiword: Optional system package for .doc file fallback
    """

    @property
    def supported_extensions(self) -> List[str]:
        """Return list of supported Word document extensions.

        Returns:
            List containing '.docx' and '.doc' extensions.
        """
        return [".docx", ".doc"]

    def process(self, file_path: Union[str, Path], **kwargs: Any) -> str:
        """Extract text content from Word documents.

        This method handles both .docx and .doc formats, extracting text from
        paragraphs and optionally from tables. It preserves document structure
        while providing clean, readable text output.

        Args:
            file_path: Path to the Word document to process.
            **kwargs: Additional processing parameters:
                - include_tables (bool): Whether to include table content
                  in the extracted text (default: True)

        Returns:
            Extracted text content with paragraphs and tables formatted as text.

        Raises:
            ProcessingError: If required dependencies are missing, file format
                           is unsupported, or processing fails.
        """
        self.validate_file(file_path)
        file_path = Path(file_path)

        if file_path.suffix.lower() == ".docx":
            return self._process_docx_file(file_path, **kwargs)
        if file_path.suffix.lower() == ".doc":
            return self._process_doc_file(file_path, **kwargs)
        raise ProcessingError(f"Unsupported file extension: {file_path.suffix}")

    def _process_docx_file(self, file_path: Path, **kwargs: Any) -> str:
        """Process .docx files using python-docx library.

        Args:
            file_path: Path to the .docx file.
            **kwargs: Processing parameters.

        Returns:
            Extracted text content.

        Raises:
            ProcessingError: If python-docx is not available or processing fails.
        """
        try:
            import docx  # pylint: disable=import-outside-toplevel
        except ImportError as exc:
            raise ProcessingError(
                "python-docx is required for .docx processing. " "Install with: pip install python-docx"
            ) from exc

        try:
            doc = docx.Document(str(file_path))
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text_content.append(paragraph_text)

            # Extract text from tables if requested
            include_tables = kwargs.get("include_tables", True)
            if include_tables:
                for table in doc.tables:
                    table_content = self._extract_table_content(table)
                    if table_content:
                        text_content.extend(table_content)

            return "\n".join(text_content)

        except Exception as e:
            raise ProcessingError(f"Failed to process .docx file {file_path}: {str(e)}") from e

    def _extract_table_content(self, table: Any) -> List[str]:
        """Extract and format content from a Word table.

        Args:
            table: python-docx Table object.

        Returns:
            List of formatted table rows as strings.
        """
        table_rows = []

        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
                else:
                    row_cells.append("")  # Preserve empty cells for structure

            if any(cell for cell in row_cells):  # Only add rows with content
                table_rows.append("\t".join(row_cells))

        return table_rows

    def _process_doc_file(self, file_path: Path, **kwargs: Any) -> str:  # pylint: disable=unused-argument
        """Process .doc files using docx2txt or antiword fallback.

        Args:
            file_path: Path to the .doc file.
            **kwargs: Processing parameters (currently unused for .doc files).

        Returns:
            Extracted text content.

        Raises:
            ProcessingError: If no suitable library/tool is available for .doc processing.
        """
        # Try docx2txt first
        try:
            import docx2txt  # pylint: disable=import-outside-toplevel

            result: str | None = docx2txt.process(str(file_path))
            return result if result is not None else ""
        except ImportError:
            pass

        # Fallback to antiword if available
        try:
            import subprocess  # pylint: disable=import-outside-toplevel

            subprocess_result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                check=True,
                timeout=30,  # Add timeout for safety
            )
            return subprocess_result.stdout
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as exc:
            raise ProcessingError(
                "Cannot process .doc files. Install python-docx2txt with: pip install docx2txt, "
                "or install antiword system package"
            ) from exc

    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract comprehensive metadata from Word documents.

        This method extracts both file system metadata and document-specific
        properties when available. Metadata extraction is more comprehensive
        for .docx files than .doc files.

        Args:
            file_path: Path to the Word document.

        Returns:
            Dictionary containing metadata with keys:
            - file_path: Original file path
            - file_size: File size in bytes
            - file_type: File extension
            - title, author, subject, keywords: Document properties (for .docx)
            - created, modified, last_modified_by: Document timestamps
            - paragraph_count, table_count, section_count: Document structure info
            - error: Error message if metadata extraction fails
        """
        self.validate_file(file_path)
        file_path = Path(file_path)

        stat = file_path.stat()
        metadata = {
            "file_path": str(file_path),
            "file_size": stat.st_size,
            "file_type": file_path.suffix.lower(),
            "created": str(stat.st_ctime),
            "modified": str(stat.st_mtime),
        }

        if file_path.suffix.lower() == ".docx":
            metadata.update(self._extract_docx_metadata(file_path))
        else:
            metadata["note"] = "Metadata extraction not supported for .doc files"

        return metadata

    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from .docx files using python-docx.

        Args:
            file_path: Path to the .docx file.

        Returns:
            Dictionary with extracted metadata.
        """
        try:
            import docx  # pylint: disable=import-outside-toplevel

            doc = docx.Document(str(file_path))

            # Extract core document properties
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "category": props.category or "",
                "comments": props.comments or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
            }

            # Count document elements for structure analysis
            metadata.update(
                {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "section_count": len(doc.sections),
                }
            )

            # Additional document analysis
            non_empty_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            metadata["non_empty_paragraph_count"] = non_empty_paragraphs

            return metadata

        except ImportError:
            return {"error": "python-docx not available for metadata extraction"}
        except Exception as e:  # pylint: disable=broad-exception-caught
            return {"error": f"Failed to extract metadata: {str(e)}"}
