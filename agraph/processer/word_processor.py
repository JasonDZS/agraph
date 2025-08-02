from pathlib import Path
from typing import Any, Dict, List, Union

from .base import DocumentProcessor, ProcessingError


class WordProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents (.docx and .doc)."""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def process(self, file_path: Union[str, Path], **kwargs: Any) -> str:
        """Extract text from Word document.

        Args:
            file_path: Path to the Word document
            **kwargs: Additional parameters

        Returns:
            Extracted text content
        """
        self.validate_file(file_path)
        file_path = Path(file_path)

        if file_path.suffix.lower() == ".docx":
            return self._process_docx(file_path, **kwargs)
        elif file_path.suffix.lower() == ".doc":
            return self._process_doc(file_path, **kwargs)
        else:
            raise ProcessingError(f"Unsupported file extension: {file_path.suffix}")

    def _process_docx(self, file_path: Path, **kwargs: Any) -> str:
        """Process .docx files using python-docx."""
        try:
            import docx
        except ImportError:
            raise ProcessingError("python-docx is required for .docx processing. Install with: pip install python-docx")

        try:
            doc = docx.Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables if requested
            include_tables = kwargs.get("include_tables", True)
            if include_tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content.append("\t".join(row_text))

            return "\n".join(text_content)

        except Exception as e:
            raise ProcessingError(f"Failed to process .docx file {file_path}: {str(e)}")

    def _process_doc(self, file_path: Path, **kwargs: Any) -> str:
        """Process .doc files using python-docx2txt or antiword fallback."""
        try:
            import docx2txt

            result: str | None = docx2txt.process(str(file_path))
            return result if result is not None else ""
        except ImportError:
            pass

        # Fallback to antiword if available
        try:
            import subprocess

            subprocess_result = subprocess.run(["antiword", str(file_path)], capture_output=True, text=True, check=True)
            return subprocess_result.stdout
        except (subprocess.CalledProcessError, FileNotFoundError):
            raise ProcessingError(
                "Cannot process .doc files. Install python-docx2txt with: pip install docx2txt, "
                "or install antiword system package"
            )

    def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract metadata from Word document.

        Args:
            file_path: Path to the Word document

        Returns:
            Dictionary containing metadata
        """
        self.validate_file(file_path)
        file_path = Path(file_path)

        metadata = {
            "file_path": str(file_path),
            "file_size": file_path.stat().st_size,
            "file_type": file_path.suffix.lower(),
        }

        if file_path.suffix.lower() == ".docx":
            try:
                import docx

                doc = docx.Document(file_path)

                # Extract core properties
                props = doc.core_properties
                metadata.update(
                    {
                        "title": props.title or "",
                        "author": props.author or "",
                        "subject": props.subject or "",
                        "keywords": props.keywords or "",
                        "category": props.category or "",
                        "comments": props.comments or "",
                        "created": str(props.created) if props.created else "",
                        "modified": str(props.modified) if props.modified else "",
                        "last_modified_by": props.last_modified_by or "",
                    }
                )

                # Count elements
                metadata.update(
                    {
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                        "section_count": len(doc.sections),
                    }
                )

            except ImportError:
                metadata["error"] = "python-docx not available for metadata extraction"
            except Exception as e:
                metadata["error"] = f"Failed to extract metadata: {str(e)}"
        else:
            metadata["note"] = "Metadata extraction not supported for .doc files"

        return metadata
